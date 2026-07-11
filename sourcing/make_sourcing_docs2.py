# -*- coding: utf-8 -*-
"""소싱 기술서 3종 v2 — 참고 판매자 링크 + 타겟 랜딩원가(공헌40%) 추가 → 로컬 docx"""
import io, os
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn

OUT = "/Users/macmini_ky/Library/CloudStorage/GoogleDrive-cky2833@gmail.com/내 드라이브/Claude AI work space/mac window file transfer/상아님_소싱_매뉴얼"
RECV, LOGI, MARGIN = 0.914, 2300, 0.40
CNY, USD = 225, 1500  # 적용환율 (2026-07-11, 보수적): ¥1=225원, $1=1,500원

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    nr = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    nr.append(rPr); t = OxmlElement('w:t'); t.text = text; nr.append(t); hl.append(nr)
    paragraph._p.append(hl)

def landing_rows(scenarios):
    rows = [["포지션","목표 판매가","타겟 랜딩원가(원)","EXW 원화","1688 목표단가(¥)","알리바바 목표단가($)"]]
    for label, price in scenarios:
        settle = price*RECV; target = settle - LOGI - price*MARGIN; src = target*0.65
        rows.append([label, f"{price:,}", f"{int(target):,}", f"~{int(src):,}", f"¥{src/CNY:.1f}", f"${src/USD:.2f}"])
    return rows

# 제품별 상위 판매자 [상품명, 판매가, 월매출표기, 리뷰, url]
SELLERS = {
"drain": [
  ["에스엠지 냄새제로 하수구트랩 50mm","14,900","월1.94억","34,582","https://www.coupang.com/vp/products/7380618078"],
  ["스멜컷 완벽차단 7세대 하수구트랩","14,300","월1.64억","2,284","https://www.coupang.com/vp/products/9073111460"],
  ["온메이드 완전밀폐 냄새제로 (프리미엄)","28,090","월1.20억","497","https://www.coupang.com/vp/products/9496868897"],
  ["에어스케이프 끝판왕 하수구트랩","24,900","월1.13억","2,682","https://www.coupang.com/vp/products/7840030377"],
  ["리빙트리 하수구 트랩 (저가)","7,090","월0.90억","5,398","https://www.coupang.com/vp/products/8643178051"],
],
"brush": [
  ["온메이드 스칼프 마사지 샴푸브러쉬 (프리미엄)","28,090","월0.80억","170","https://www.coupang.com/vp/products/9497239908"],
  ["다슈 데일리 스칼프 스케일링 브러쉬","8,450","월0.48억","16,693","https://www.coupang.com/vp/products/7153349754"],
  ["라보에이치 프리미엄 샴푸브러쉬 (1위)","6,910","월0.45억","5,151","https://www.coupang.com/vp/products/5302545669"],
  ["쿤달 스칼프 마사지 샴푸브러쉬","11,110","월0.41억","10,028","https://www.coupang.com/vp/products/6806892955"],
  ["로렌코스 프로페셔널 두피 브러쉬","8,900","월0.25억","6,079","https://www.coupang.com/vp/products/5102692658"],
],
"band": [
  ["써머텍트 기능성 헤어밴드 (1위)","11,800","월1.42억","16,480","https://www.coupang.com/vp/products/5716566331"],
  ["아이더 세이프티 러닝 헤어밴드 (브랜드)","24,340","월1.27억","518","https://www.coupang.com/vp/products/9496716071"],
  ["써머텍트 기능성 2개입","8,850","월1.02억","16,480","https://www.coupang.com/vp/products/5716566331"],
  ["나이키 스우시 클래식 헤어밴드 2P","23,990","월0.45억","318","https://www.coupang.com/vp/products/8935438600"],
  ["딥루트 스포츠 땀흡수 헤어밴드 (저가)","7,610","월0.39억","590","https://www.coupang.com/vp/products/9384723915"],
],
}
LANDING = {
"drain": [("프리미엄",19900),("주력",14900),("진입",9900)],
"brush": [("중가",12900),("주력",8900),("진입",6900)],
"band":  [("세트 중가",12900),("주력",9900),("진입",7900)],
}

# ===== 문서 정의 (v1 본문 + 신규 두 섹션) =====
def common_tail(key, extra_note=""):
    return [
      ("★ 참고 상위 판매자 (실물 벤치마킹 · 클릭)", ["sellers", key]),
      ("★ 타겟 랜딩원가 역산 (공헌이익 40% 기준)", ["landing", key]),
      ("   원가 산식 안내", ["para",
        "정산수취 = 판매가×0.914(쿠팡수수료·쿠폰 차감, 그로스 6월 실정산값) / 물류비 2,300원/건(소형, 실부과 60%·2027.1.31 저가할인 종료 리스크) / 마진40%=광고재원+순이익 / EXW 원화=랜딩원가×0.65(관세·물류·검수 제외). 1688은 위안(¥), 알리바바는 달러($) 결제 통화로 환산 — 적용환율 ¥1=225원·$1=1,500원(2026-07-11 기준, 발주 시 재확인). "+extra_note]),
    ]

# v1 본문 재사용
import importlib.util
spec = importlib.util.spec_from_file_location("v1", "/Users/macmini_ky/ClaudeAITeam/sourcing/make_sourcing_docs.py")
v1 = importlib.util.module_from_spec(spec)
_src = open("/Users/macmini_ky/ClaudeAITeam/sourcing/make_sourcing_docs.py").read().split("def main()")[0]
_ns = {}; exec(_src, _ns)
V1DOCS = {d["title"]: d for d in _ns["DOCS"]}

KEYMAP = {"[소싱] 배수구 트랩 기술서":"drain", "[소싱] 샴푸 브러쉬 기술서":"brush", "[소싱] 남자 헤어밴드 기술서":"band"}
EXTRA = {"drain":"", "brush":"", "band":"헤어밴드는 4~6개입 세트 기준 총 랜딩원가 → 개당은 구성수로 나눔(예: 4개입이면 ÷4)."}

def build(doc_data, key):
    d = Document()
    d.add_heading(doc_data["title"], level=0)
    p = d.add_paragraph(); r = p.add_run(doc_data["sub"]); r.italic = True; r.font.size = Pt(9)
    sections = list(doc_data["sections"])
    # '타겟 가격·원가' 섹션 뒤에 신규 삽입: 맨 끝 TIP 앞에
    tip_idx = next((i for i,(n,_) in enumerate(sections) if n=="상아님 TIP"), len(sections))
    sections = sections[:tip_idx] + common_tail(key, EXTRA[key]) + sections[tip_idx:]
    for name, body in sections:
        d.add_heading(name, level=1)
        kind = body[0]
        if kind == "para": d.add_paragraph(body[1])
        elif kind == "bullets":
            for it in body[1:]: d.add_paragraph(it, style="List Bullet")
        elif kind == "numbers":
            for it in body[1:]: d.add_paragraph(it, style="List Number")
        elif kind == "table":
            rows = body[1:]
            t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    t.cell(ri, ci).text = val
                    if ri == 0:
                        for pp in t.cell(ri,ci).paragraphs:
                            for rr in pp.runs: rr.font.bold = True
        elif kind == "landing":
            rows = landing_rows(LANDING[body[1]])
            t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    t.cell(ri, ci).text = val
                    if ri == 0:
                        for pp in t.cell(ri,ci).paragraphs:
                            for rr in pp.runs: rr.font.bold = True
        elif kind == "sellers":
            data = SELLERS[body[1]]
            t = d.add_table(rows=len(data)+1, cols=5); t.style = "Light Grid Accent 1"
            hdr = ["상품 (링크)","판매가","월매출","리뷰수","쿠팡 링크"]
            for ci,h in enumerate(hdr):
                t.cell(0,ci).text=h
                for pp in t.cell(0,ci).paragraphs:
                    for rr in pp.runs: rr.font.bold=True
            for ri,(nm,pr,rev,rv,url) in enumerate(data, start=1):
                t.cell(ri,0).text=nm; t.cell(ri,1).text=pr; t.cell(ri,2).text=rev; t.cell(ri,3).text=rv
                cell=t.cell(ri,4); cell.text=""
                add_hyperlink(cell.paragraphs[0], url, "바로가기")
    path = os.path.join(OUT, doc_data["title"]+".docx")
    d.save(path); return path, len(d.tables)

for title, key in KEYMAP.items():
    path, ntab = build(V1DOCS[title], key)
    print(f"OK | {os.path.basename(path)} | 표 {ntab}개 | {os.path.getsize(path):,} bytes")
